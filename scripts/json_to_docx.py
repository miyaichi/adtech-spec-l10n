#!/usr/bin/env python3
"""
JSON to DOCX Converter

翻訳済みJSONファイルからWord文書(DOCX)を生成します。
対訳モード（英語+日本語）と日本語のみモードをサポートします。
"""
import json
import sys
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def create_docx(json_data, output_file, document_title=None, bilingual_mode=False,
                font_name='Yu Mincho', font_size=10.5):
    """
    JSONデータからDOCX文書を生成します。

    Args:
        json_data: 翻訳データ（辞書のリスト）
        output_file: 出力ファイルパス
        document_title: ドキュメントタイトル
        bilingual_mode: 対訳モードの有効/無効
        font_name: フォント名
        font_size: フォントサイズ
    """
    doc = Document()

    # スタイル設定
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)

    # ドキュメントタイトル
    if document_title:
        doc.add_heading(document_title, 0)
    else:
        doc.add_heading('Document Translation', 0)

    last_section = None

    for entry in json_data:
        section = entry.get('section', '')
        en_text = entry.get('en', '')
        ja_text = entry.get('ja', '')

        # セクションが変わったら見出しを追加
        if section != last_section and section:
            # 見出しレベルの判定（数字の階層に基づく）
            # "1. ABSTRACT" -> Level 1
            # "3.1 ACCESS METHOD" -> Level 2
            level = min(section.count('.') + 1, 3)

            # 見出しを追加
            doc.add_heading(section, level=level)
            last_section = section

        # 本文の追加
        if bilingual_mode:
            # 対訳モード: 英語（小さめ） -> 日本語
            if en_text:
                p_en = doc.add_paragraph(en_text)
                if p_en.runs:
                    p_en.runs[0].font.size = Pt(9)
                    p_en.runs[0].font.color.rgb = RGBColor(100, 100, 100)  # グレー

            if ja_text:
                p_ja = doc.add_paragraph(ja_text)
                p_ja_fmt = p_ja.paragraph_format
                p_ja_fmt.space_after = Pt(12)  # 段落後のスペース
            else:
                # 翻訳がない場合は警告を表示
                p = doc.add_paragraph("[未翻訳]")
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 赤
        else:
            # 日本語のみモード
            if ja_text:
                doc.add_paragraph(ja_text)
            else:
                # 訳文がない場合は原文を表示して警告
                p = doc.add_paragraph(f"[TODO: 翻訳が必要] {en_text}")
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(255, 100, 0)  # オレンジ

    doc.save(output_file)
    print(f"Successfully saved to {output_file}")


def main():
    parser = argparse.ArgumentParser(
        description='翻訳済みJSONからDOCX文書を生成します',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用例:
  # 日本語のみモード
  python3 json_to_docx.py translated.json output.docx

  # 対訳モード（英語+日本語）
  python3 json_to_docx.py translated.json output.docx --bilingual

  # タイトルとフォントを指定
  python3 json_to_docx.py translated.json output.docx \\
      --title "Ads.txt Specification v1.1 (Japanese)" \\
      --font "Hiragino Mincho ProN" --font-size 11
        '''
    )

    parser.add_argument('input', help='入力JSON（翻訳済み）ファイル')
    parser.add_argument('output', help='出力DOCX ファイル')
    parser.add_argument('--bilingual', action='store_true',
                       help='対訳モード（英語+日本語を表示）')
    parser.add_argument('--title', help='ドキュメントのタイトル')
    parser.add_argument('--font', default='Yu Mincho',
                       help='フォント名（デフォルト: Yu Mincho）')
    parser.add_argument('--font-size', type=float, default=10.5,
                       help='フォントサイズ（デフォルト: 10.5）')

    args = parser.parse_args()

    try:
        # 入力ファイルの読み込み
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"Error: Input file not found: {args.input}", file=sys.stderr)
            sys.exit(1)

        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # 出力ディレクトリの作成
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # DOCX生成
        create_docx(
            data,
            args.output,
            document_title=args.title,
            bilingual_mode=args.bilingual,
            font_name=args.font,
            font_size=args.font_size
        )

    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON format: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
